"""Per-format text extraction.

Accepts *any* file: rich formats (PDF, DOCX, PPTX, XLSX, CSV, images) have
dedicated handlers; everything else (TXT, MD, JSON, XML, YAML, HTML, source
code, logs, …) goes through a universal text decoder with a binary guard, so
unknown-but-textual files still ingest and true binaries fail with a clear error.
"""
from __future__ import annotations

import csv
import io

from app.core.exceptions import ProcessingError, UnsupportedFileTypeError
from app.services.document_processing.ocr import OCREngine

# Extensions with dedicated rich extractors. Anything not listed is treated as
# generic text (handled by the universal decoder below).
RICH_TYPES: dict[str, str] = {
    "pdf": "pdf",
    "docx": "docx",
    "pptx": "pptx",
    "xlsx": "xlsx",
    "xls": "xlsx",
    "csv": "csv",
    "html": "html",
    "htm": "html",
    "png": "image",
    "jpg": "image",
    "jpeg": "image",
    "webp": "image",
    "bmp": "image",
    "tif": "image",
    "tiff": "image",
    "gif": "image",
}

# Listed only for documentation/UX; extraction itself accepts any extension.
TEXTUAL_TYPES = {
    "txt", "md", "markdown", "json", "xml", "yaml", "yml", "toml", "ini",
    "log", "csv", "tsv", "rst", "py", "js", "jsx", "ts", "tsx", "java", "go",
    "rb", "rs", "c", "cpp", "h", "sh", "sql", "env", "cfg", "conf",
}


def detect_file_type(file_name: str) -> str:
    """Return the handler tag for a file name. Unknown extensions -> 'text'.

    Never rejects by extension: binary detection happens during extraction so
    the platform can accept arbitrary uploads and fail gracefully on true binaries.
    """
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    return RICH_TYPES.get(ext, "text")


def extract(data: bytes, file_type: str, ocr: OCREngine) -> tuple[str, int]:
    """Extract text from raw bytes. Returns (text, page_count)."""
    try:
        if file_type == "pdf":
            return _extract_pdf(data)
        if file_type == "docx":
            return _extract_docx(data), 1
        if file_type == "pptx":
            return _extract_pptx(data)
        if file_type == "xlsx":
            return _extract_xlsx(data), 1
        if file_type == "csv":
            return _extract_csv(data), 1
        if file_type == "html":
            return _extract_html(data), 1
        if file_type == "image":
            return ocr.image_to_text(data), 1
        # Universal fallback: decode as text with a binary guard.
        return _decode_text(data), 1
    except ProcessingError:
        raise
    except UnsupportedFileTypeError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ProcessingError(f"Failed to extract text: {exc}") from exc


def _extract_pdf(data: bytes) -> tuple[str, int]:
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        page_count = doc.page_count
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts), page_count


def _extract_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pptx(data: bytes) -> tuple[str, int]:
    """Pull slide text from a .pptx (OOXML zip) using only the stdlib."""
    import zipfile
    from xml.etree import ElementTree as ET

    text_tag = "{http://schemas.openxmlformats.org/drawingml/2006/main}t"
    parts: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        slides = sorted(
            n for n in zf.namelist()
            if n.startswith("ppt/slides/slide") and n.endswith(".xml")
        )
        for name in slides:
            root = ET.fromstring(zf.read(name))
            runs = [el.text for el in root.iter(text_tag) if el.text]
            if runs:
                parts.append(" ".join(runs))
    return "\n".join(parts), max(1, len(parts))


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(cells):
                parts.append(", ".join(cells))
    wb.close()
    return "\n".join(parts)


def _extract_csv(data: bytes) -> str:
    text = _decode_text(data)
    reader = csv.reader(io.StringIO(text))
    return "\n".join(", ".join(row) for row in reader)


def _extract_html(data: bytes) -> str:
    from html.parser import HTMLParser

    class _Stripper(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self.chunks: list[str] = []
            self._skip = False

        def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ARG002
            if tag in ("script", "style"):
                self._skip = True

        def handle_endtag(self, tag: str) -> None:
            if tag in ("script", "style"):
                self._skip = False

        def handle_data(self, data: str) -> None:
            if not self._skip and data.strip():
                self.chunks.append(data.strip())

    parser = _Stripper()
    parser.feed(data.decode("utf-8", errors="replace"))
    return "\n".join(parser.chunks)


def _decode_text(data: bytes) -> str:
    """Decode bytes as text, rejecting binary content."""
    if not data:
        raise ProcessingError("File is empty")

    # A NUL byte is a strong binary signal — text formats never contain one.
    if b"\x00" in data[:8192]:
        raise UnsupportedFileTypeError(
            "File appears to be binary or in an unsupported format"
        )

    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        decoded = data.decode("utf-8", errors="replace")
        sample = decoded[:4096]
        # Many undecodable bytes -> not real text.
        if sample.count("�") / max(1, len(sample)) > 0.1:
            raise UnsupportedFileTypeError(
                "File appears to be binary or in an unsupported format"
            )
        text = data.decode("latin-1", errors="replace")

    if not text.strip():
        raise ProcessingError("No readable text in file")
    return text
